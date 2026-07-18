from types import SimpleNamespace
from io import BytesIO

from docx import Document

from travel_agent.services.upload_service import extract_upload_context


def test_extract_txt_upload():
    upload = SimpleNamespace(filename='note.txt', content_type='text/plain')
    context = extract_upload_context(upload, '你好\n世界'.encode('utf-8'))
    assert context['file_kind'] == 'text'
    assert '你好' in context['extracted_text']


def test_extract_markdown_upload():
    upload = SimpleNamespace(filename='plan.md', content_type='text/markdown')
    context = extract_upload_context(upload, b'# Title\n\nbody')
    assert context['file_kind'] == 'text'
    assert 'Title' in context['extracted_text']


def test_extract_docx_upload_reads_paragraphs_and_tables():
    document = Document()
    document.add_paragraph('旅行需求：杭州三天两晚')
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = '偏好'
    table.cell(0, 1).text = '博物馆'
    buffer = BytesIO()
    document.save(buffer)

    upload = SimpleNamespace(
        filename='plan.docx',
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    context = extract_upload_context(upload, buffer.getvalue())

    assert context['file_kind'] == 'document'
    assert '杭州三天两晚' in context['extracted_text']
    assert '博物馆' in context['extracted_text']


def test_image_upload_is_honest_without_content_recognition():
    upload = SimpleNamespace(filename='photo.png', content_type='image/png')
    context = extract_upload_context(upload, b'\x89PNG\r\n')
    assert context['file_kind'] == 'image'
    assert context['extracted_text'] == ''
    assert context['extraction_error'] == 'image_content_not_parsed'
    assert '暂未解析图片内容' in context['notice']
