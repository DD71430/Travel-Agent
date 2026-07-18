from __future__ import annotations

from io import BytesIO
import mimetypes
import re
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

from travel_agent.core.config import get_settings
from travel_agent.core.logging import get_logger
from travel_agent.services.audio_service import transcribe_audio

settings = get_settings()
logger = get_logger(__name__)

MAX_UPLOAD_BYTES = settings.max_upload_bytes
MAX_EXTRACT_CHARS = settings.max_extract_chars
TEXT_EXTENSIONS = {'.txt', '.md', '.csv', '.log', '.json'}
AUDIO_EXTENSIONS = {'.mp3', '.wav', '.m4a', '.aac', '.ogg', '.webm', '.flac', '.opus', '.amr', '.3gp'}
DOCX_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def truncate_text(text: str, limit: int = MAX_EXTRACT_CHARS) -> str:
    compact = text.strip()
    return compact if len(compact) <= limit else compact[:limit]


def normalize_extracted_text(text: str) -> str:
    lines: list[str] = []
    previous = ''
    for raw_line in text.splitlines():
        line = re.sub(r'\s+', ' ', raw_line).strip()
        if not line:
            if lines and lines[-1] != '':
                lines.append('')
            continue
        if line == previous or (len(line) <= 2 and previous and previous.endswith(line)):
            continue
        lines.append(line)
        previous = line
    return truncate_text(re.sub(r'\n{3,}', '\n\n', '\n'.join(lines)))


def _looks_like_pdf_scan(content: bytes) -> bool:
    if b'/Font' in content or b'/Text' in content:
        return False
    return content.count(b'\x00') < max(20, len(content) // 50)


def extract_text_from_pdf(content: bytes) -> str:
    extracted_parts: list[str] = []
    try:
        reader = PdfReader(BytesIO(content))
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ''
            except Exception:
                logger.debug('PDF page text extraction failed', exc_info=True)
                page_text = ''
            if page_text.strip():
                extracted_parts.append(page_text)
    except Exception:
        logger.exception('PDF extraction failed')
        extracted_parts = []
    text = '\n'.join(extracted_parts).strip()
    if text:
        return normalize_extracted_text(text)
    if _looks_like_pdf_scan(content):
        return ''
    return ''


def extract_text_from_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(' '.join(cells))
    return normalize_extracted_text('\n'.join(parts))


def _decode_plain_bytes(content: bytes) -> str:
    for encoding in ('utf-8-sig', 'utf-8', 'gb18030', 'gbk', 'big5', 'latin-1'):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return ''


def extract_text_from_plain_file(content: bytes) -> str:
    decoded = _decode_plain_bytes(content)
    if not decoded:
        return ''
    decoded = decoded.replace('\r\n', '\n').replace('\r', '\n')
    decoded = re.sub(r'\u0000+', '', decoded)
    decoded = re.sub(r'(?m)^[\s\-_=\*#]{3,}$', '', decoded)
    decoded = re.sub(r'(?m)^\s*\d+\s*$', '', decoded)
    return normalize_extracted_text(decoded)


def extract_upload_context(upload: UploadFile | None, content: bytes | None) -> dict | None:
    if upload is None or content is None:
        return None
    filename = upload.filename or 'uploaded-file'
    suffix = Path(filename).suffix.lower()
    content_type = upload.content_type or mimetypes.guess_type(filename)[0] or 'application/octet-stream'
    if len(content) > MAX_UPLOAD_BYTES:
        return {'filename': filename, 'content_type': content_type, 'size': len(content), 'file_kind': 'unknown', 'extracted_text': '', 'extraction_error': 'file_too_large'}
    extracted_text = ''
    file_kind = 'binary'
    extraction_error = ''
    audio_debug: dict = {}
    if content_type.startswith('image/'):
        file_kind = 'image'
        extraction_error = 'image_content_not_parsed'
    elif suffix == '.pdf' or content_type == 'application/pdf':
        file_kind = 'pdf'
        try:
            extracted_text = extract_text_from_pdf(content)
        except Exception:
            logger.exception('PDF upload extraction failed')
            extraction_error = 'pdf_extract_failed'
    elif suffix == '.docx' or content_type == DOCX_CONTENT_TYPE:
        file_kind = 'document'
        try:
            extracted_text = extract_text_from_docx(content)
        except Exception:
            logger.exception('DOCX upload extraction failed')
            extraction_error = 'docx_extract_failed'
    elif suffix in TEXT_EXTENSIONS or content_type.startswith('text/'):
        file_kind = 'text'
        try:
            extracted_text = extract_text_from_plain_file(content)
        except Exception:
            logger.exception('Text upload extraction failed')
            extraction_error = 'text_extract_failed'
    elif suffix in AUDIO_EXTENSIONS or content_type.startswith('audio/'):
        file_kind = 'audio'
        extracted_text, audio_error, audio_debug = transcribe_audio(filename, content, content_type)
        if not extracted_text:
            extraction_error = audio_error or 'audio_transcribe_failed'
    upload_context = {'filename': filename, 'content_type': content_type, 'size': len(content), 'file_kind': file_kind, 'extracted_text': extracted_text, 'extraction_error': extraction_error}
    if file_kind == 'image':
        upload_context['notice'] = '已接收图片；当前版本暂未解析图片内容，可补充文字需求。'
    if audio_debug and settings.debug:
        upload_context['audio_debug'] = audio_debug
    return upload_context


def merge_question_with_file_context(question: str, file_context: dict | None) -> str:
    if not file_context:
        return question
    extracted_text = str(file_context.get('extracted_text') or '').strip()
    if extracted_text:
        return extracted_text
    error_text = str(file_context.get('extraction_error') or '').strip()
    return question if question.strip() else error_text
